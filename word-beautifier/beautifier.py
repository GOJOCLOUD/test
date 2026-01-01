# beautifier.py  —— Word 文档美化核心逻辑（强制区分大标题 / 正文，并统一排版）

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from tqdm import tqdm


def _set_run_font(run, font_name, font_size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold

    # 兼容中文字体
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), font_name)


def _format_paragraph(
    para,
    font_name,
    font_size,
    line_spacing,
    space_before,
    space_after,
    first_line_indent_cm,
    align,
    bold=False,
):
    # 文字
    if not para.runs:
        run = para.add_run("")
        _set_run_font(run, font_name, font_size, bold=bold)
    else:
        for run in para.runs:
            _set_run_font(run, font_name, font_size, bold=bold)

    # 段落
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_cm > 0:
        pf.first_line_indent = Cm(first_line_indent_cm)
    else:
        pf.first_line_indent = Cm(0)
    para.alignment = align


def _is_heading(text, para):
    t = text.strip()
    if not t:
        return False

    # 明确样式
    if para.style and para.style.name and para.style.name.lower().startswith("heading"):
        return True

    # 行短
    if len(t) <= 25:
        return True

    # 大写比例
    letters = [c for c in t if c.isalpha()]
    if letters:
        upper_ratio = sum(c.isupper() for c in letters) / len(letters)
        if upper_ratio > 0.4:
            return True

    # 加粗
    if any(run.bold for run in para.runs):
        return True

    return False


def beautify_document(input_path, output_path, config):
    doc = Document(input_path)

    font_name = config.get("font", "宋体")
    body_size = config.get("size", 12)
    line_spacing = config.get("line_spacing", 1.5)
    space_before = config.get("space_before", 6)
    space_after = config.get("space_after", 6)

    heading_size = body_size + 6
    heading_space_before = 18
    heading_space_after = 12

    # 正文首行缩进（2 字左右）
    body_first_line_indent_cm = 0.74

    # 1. 普通段落
    for para in tqdm(doc.paragraphs, desc="格式化正文与标题", ncols=80):
        text = para.text.strip()

        # 全空行：统一清理成「空段落但保留间距」
        if not text:
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            pf.space_before = Pt(0)
            pf.space_after = Pt(space_after)
            pf.first_line_indent = Cm(0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if _is_heading(text, para):
            # 大标题：居中 + 大号 + 加粗 + 上下留白
            _format_paragraph(
                para,
                font_name=font_name,
                font_size=heading_size,
                line_spacing=1.5,
                space_before=heading_space_before,
                space_after=heading_space_after,
                first_line_indent_cm=0,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
            )
        else:
            # 正文：统一宋体 12pt，1.5 倍行距，两端对齐，首行缩进
            _format_paragraph(
                para,
                font_name=font_name,
                font_size=body_size,
                line_spacing=line_spacing,
                space_before=space_before,
                space_after=space_after,
                first_line_indent_cm=body_first_line_indent_cm,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                bold=False,
            )

    # 2. 表格内部文字（全部按正文处理，不做标题判断）
    for table in tqdm(doc.tables, desc="格式化表格内容", ncols=80):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    _format_paragraph(
                        para,
                        font_name=font_name,
                        font_size=body_size,
                        line_spacing=line_spacing,
                        space_before=0,
                        space_after=0,
                        first_line_indent_cm=0,
                        align=WD_ALIGN_PARAGRAPH.LEFT,
                        bold=False,
                    )

    doc.save(output_path)
