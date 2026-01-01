#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
论文排版工具
功能：读取Word文档，进行排版优化，输出格式化后的新文档
"""

import os
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def format_paper(input_path, output_path):
    """
    格式化论文Word文档
    
    Args:
        input_path: 输入Word文档路径
        output_path: 输出Word文档路径
    
    Returns:
        bool: 是否成功
    """
    try:
        # 1. 打开输入文档
        doc = Document(input_path)
        
        # 2. 创建新文档用于输出
        output_doc = Document()
        
        # 3. 添加样式
        _add_paper_styles(output_doc)
        
        # 4. 处理文档内容
        _process_document_content(doc, output_doc)
        
        # 5. 保存输出文档
        output_doc.save(output_path)
        
        print(f"✅ 论文格式化成功！输出文件: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 论文格式化失败: {str(e)}")
        return False

def _add_paper_styles(doc):
    """添加论文所需的样式"""
    # 获取样式集合
    styles = doc.styles
    
    # 添加标题样式
    try:
        title_style = styles.add_style('Paper Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(18)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        
        title_para_format = title_style.paragraph_format
        title_para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para_format.space_after = Pt(24)
    except:
        pass  # 忽略已存在样式的错误
    
    # 添加一级标题样式
    try:
        heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(16)
        heading1_font.bold = True
        
        heading1_para_format = heading1_style.paragraph_format
        heading1_para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1_para_format.space_before = Pt(12)
        heading1_para_format.space_after = Pt(6)
    except:
        pass
    
    # 添加正文样式
    try:
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        
        normal_para_format = normal_style.paragraph_format
        normal_para_format.space_after = Pt(6)
        normal_para_format.line_spacing = 1.5
        normal_para_format.left_indent = Inches(0.5)
    except:
        pass

def _process_document_content(input_doc, output_doc):
    """处理文档内容"""
    # 简单处理逻辑，实际应用中可以根据需要扩展
    for para in input_doc.paragraphs:
        # 获取段落文本
        text = para.text.strip()
        
        if not text:
            # 保留空行
            output_doc.add_paragraph()
            continue
        
        # 判断段落类型并应用相应样式
        if _is_title(text):
            # 添加标题
            title_para = output_doc.add_paragraph(text, style='Paper Title')
        elif _is_heading_level_1(text):
            # 添加一级标题
            heading_para = output_doc.add_paragraph(text, style='Heading 1')
        else:
            # 添加正文
            body_para = output_doc.add_paragraph(text)
            # 设置段落格式
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def _is_title(text):
    """判断是否为标题"""
    # 简单判断逻辑：如果文本长度适中且全大写或有特殊标记
    return len(text) < 100 and (text.isupper() or '【标题】' in text or 'TITLE' in text)

def _is_heading_level_1(text):
    """判断是否为一级标题"""
    # 简单判断逻辑：以数字或特殊符号开头
    return text.startswith(('1.', '一、', '第一章', '第1章', 'Section 1', '1 '))

# 添加用于从Node.js调用的函数
def format_paper_from_js(input_file, output_file):
    """
    供JavaScript调用的格式化函数
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径
    
    Returns:
        dict: 包含结果信息的字典
    """
    try:
        success = format_paper(input_file, output_file)
        return {
            'success': success,
            'message': '论文排版成功' if success else '论文排版失败',
            'output_file': output_file if success else None
        }
    except Exception as e:
        return {
            'success': False,
            'message': f'处理异常: {str(e)}',
            'output_file': None
        }

if __name__ == '__main__':
    # 命令行参数解析
    parser = argparse.ArgumentParser(description='论文Word文档排版工具')
    parser.add_argument('--input', '-i', required=True, help='输入Word文档路径')
    parser.add_argument('--output', '-o', required=True, help='输出Word文档路径')
    
    args = parser.parse_args()
    
    # 执行格式化
    format_paper(args.input, args.output)